pip3 install python-docx
import requests
import time
from selenium import webdriver
from bs4 import BeautifulSoup
from selenium.webdriver.common.keys import Keys

import docx
import os

os.chdir('C:/crawler') #name of the directory where target file is located 
doc = docx.Document(python.docx) #python.docx is name of the word doc

#This navigates and browses the entire site
#call gig_extract to into a particular freelancer and extract data in it.

browser=webdriver.Chrome('C:\chromedriver.exe')

def crawler():

    url="https://www.fiverr.com"
    #browser=webdriver.Chrome('C:\chromedriver.exe')
    browser.get(url)
    soup=BeautifulSoup(browser.page_source)
    
    data=soup.find_all('div',attrs={'class': 'menu-cont'});

    for div in data:
            data1=div.find_all('ul');
            
            for ul in data1:
                links=ul.find_all('a')
                
                for a in links:
                    url="https://www.fiverr.com"+ a['href']+"#layout=auto&page=1"
                    browser.get(url)
                    time.sleep(20)

                    browser.find_element_by_class_name('gig-load-more').click()
                    for i in range(1,8000):
                        browser.execute_script("window.scrollTo(0, document.body.scrollHeight);")

                    count=0
                    soup=BeautifulSoup(browser.page_source)
                    
                    for link in soup.find_all('a',{'class': 'gig-link-main'}):
                        
                        count+=1
                        
                        if(count<=1):
                            url1="https://fiverr.com"+(link.get('href'))
                            browser.get(url1)

                            time.sleep(20)
                            i=1
                            
                            while(i>0):
                                try:
                                    browser.find_element_by_link_text("Show More").click()
                                except:
                                    break
                            
                            for i in range(1,1000):
                                browser.execute_script("window.scrollTo(0, document.body.scrollHeight);")

                            soup=BeautifulSoup(browser.page_source)                    
                            
                            #################################################################
                            
                            doc.add_heading('Gig no.'+i)
                            
                            for h4 in soup.find_all('h4',{'class': 'gig-fancy-header'}):
                                print(h4.text)
                                doc.add_paragraph(h4.text, style='List Bullet')
                                
                            
                            for span in soup.find_all('span',{'class': 'numeric-rating'}):
                                print("Rating: "+ span.text)
                                doc.add_paragraph('Rating: '+ span.text, style='List Bullet')
                            
                            for div in soup.find_all('div',{'class': 'gig-main-desc'}):
                                print(div.text)
                                doc.add_paragraph(div.text, style='List Bullet')
                            
                            sum=0
                            
                            data2=soup.find_all('ul',{'class', 'reviews-list'});
                            
                            for ul in data2:
                                for li in ul.find_all('li'):
                                    
                                    linkss=li.find_all('a')
                                    
                                    for a in linkss:
                                        print(a.text)
                                        doc.add_paragraph(a.text, style='List Bullet')
                                    
                                    divi=li.find_all('div',{'class', 'msg-body'});
                                    
                                    for div in divi:
                                        print(div.text)
                                        doc.add_paragraph(div.text, style='List Bullet')
                                    
                                    spann=li.find_all('span',{'class', 'rating-date'});
                                    
                                    for span in spann:
                                        print(span.text)
                                        doc.add_paragraph(span.text, style='List Bullet')
                                    
                                    sum+=1
                                    time.sleep(1)
                            
                            print(sum)
                            doc.add_paragraph('total reviews : '+sum, style='List Bullet')
                        else : break


crawler()
