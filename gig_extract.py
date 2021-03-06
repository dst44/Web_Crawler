pip3 install python-docx

import requests
import time
import select

from selenium import webdriver
from bs4 import BeautifulSoup


import docx
import os

os.chdir('C:/crawler') #name of the directory where target file is located 
doc = docx.Document(python2.docx) #python.docx is name of the word doc

#This extracts all the data of a particular freelancer:
#rating, rating date, reviewer's name and the review itself.

from selenium.webdriver.common.keys import Keys

def crawler():
    url="https://www.fiverr.com/ing_model/hold-a-sign-with-your-message?context=adv.cat_4.subcat_1&context_type=rating&pos=1&funnel=29275110-6e24-4b22-aad1-028abe6fa221"
    browser=webdriver.Chrome('C:\chromedriver.exe')
    browser.get(url)
    
    i=1
    while(i<100):
        try:
            browser.find_element_by_link_text('Show More').click()
            i+=1
        except:
            i+=1

    soup=BeautifulSoup(browser.page_source)
    
    print("{")
    
    for d1 in soup.find_all('span',{'class': 'gig-title'}):
        print('"'+"Gig_name"+'"'+": "+'"'+d1.text+'",')
        doc.add_heading("Gig_name: "+d1.text)
        
    #for d2 in soup.find_all()
    
    for h4 in soup.find_all('h4',{'class': 'gig-fancy-header'}):
        #for s in h4.split() if s.isdigit()]
        
        for s in h4.text.split():
            if s.isdigit():
                num=int(s)
                #l = str()
                #print(l)
                print('"'+"Review-count"+'"'+": "+'"'+str(num)+'",')
                doc.add_paragraph("Review Count: "+str(num), style='List Bullet')
    
    for span in soup.find_all('span',{'class': 'numeric-rating'}):
        print('"'+"Rating"+'"'+": "+'"'+span.text+'",')
        doc.add_paragraph("Rating: " + span.text, style='List Bullet')
    
    for div in soup.find_all('div',{'class': 'gig-main-desc'}):
         print('"'+'about_the_gig'+'"'+':'+'"'+div.text+'",')
         doc.add_paragraph("about_the_gig: "+div.text, style='List Bullet')
       # print(div.text)
         sum=0
         print('"'+"Reviews"+'"'+': ')
         print("[")
         doc.add_heading("Reviews")
    
    data2=soup.find_all('ul',{'class':'reviews-list'});
    
    for ul in data2:
        z=0
        
        for li in ul.find_all('li'):
            if z==0:
                print("{")
            linkss=li.find_all('a')
            
            for a in linkss:
                print('"'+"reviewer_name"+'"'+": "+'"'+a.text+'",')
                doc.add_paragraph("reviewer_name: "+a.text, style='List Bullet')
            
            divi=li.find_all('div',{'class': 'msg-body'});
            
            for div in divi:
                if z==0:
                    print('"'+"text"+'"'+": "+'"'+div.text+'",')
                    doc.add_paragraph("text: "+div.text, style='List Bullet')
                    z=1
                else:
                    print('"'+"Feedback"+'"'+": "+'"'+div.text+'",')
                    doc.add_paragraph("Feedback: "+div.text, style='List Bullet')
                    z=0
            
            spann=li.find_all('span',{'class': 'rating-date'});
            
            for span in spann:
                print('"'+"time"+'"'+": "+'"'+span.text+'",')
                doc.add_paragraph("time: "+span.text, style='List Bullet')
            
            sum+=1
            
            if z==0:
                print("}")

    print("]")
    print("}")

crawler()
